"""Document processing (PDF/DOCX generation)."""

from __future__ import annotations

import base64
import io
from dataclasses import dataclass

from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document as DocxDocument

# NOTE (Windows): WeasyPrint requires native GTK/Pango dependencies.
# To keep the service runnable cross-platform, we lazy-import WeasyPrint and
# fall back to ReportLab for PDF generation when WeasyPrint isn't available.


@dataclass(frozen=True)
class GeneratedDocument:
    bytes_data: bytes
    content_type: str
    output_format: str  # "pdf" | "docx"
    embedded_images: int = 0
    linked_images: int = 0


class DocumentProcessor:
    """Processing layer component: generate documents from filtered HTML.

    Rules:
    - Must include source attribution (handled by adding a header section)
    - Must support PDF and DOCX
    - Must not perform storage operations
    """

    def generate_pdf(
        self,
        filtered_html: str,
        base_url: str,
        source_url: str,
        *,
        embedded_images: int = 0,
        linked_images: int = 0,
        reportlab_images: list[tuple[str, bytes, str, str, str]] | None = None,
    ) -> GeneratedDocument:
        # Try WeasyPrint first (best HTML fidelity), then fall back to ReportLab.
        pdf_bytes = None
        html = self._wrap_with_attribution(filtered_html, source_url)

        try:
            from weasyprint import HTML  # type: ignore

            pdf_bytes = HTML(string=html, base_url=base_url).write_pdf()
        except Exception:
            pdf_bytes = self._generate_pdf_reportlab(filtered_html, source_url, reportlab_images or [])

        return GeneratedDocument(
            bytes_data=pdf_bytes,
            content_type="application/pdf",
            output_format="pdf",
            embedded_images=embedded_images,
            linked_images=linked_images,
        )

    def generate_docx(
        self,
        filtered_html: str,
        source_url: str,
        *,
        embedded_images: int = 0,
        linked_images: int = 0,
    ) -> GeneratedDocument:
        doc = DocxDocument()
        doc.add_paragraph(f"Source: {source_url}")
        doc.add_paragraph("")

        soup = BeautifulSoup(filtered_html, "lxml")
        text = soup.get_text("\n", strip=True)
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())

        buf = io.BytesIO()
        doc.save(buf)
        return GeneratedDocument(
            bytes_data=buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            output_format="docx",
            embedded_images=embedded_images,
            linked_images=linked_images,
        )

    def _wrap_with_attribution(self, html: str, source_url: str) -> str:
        return f"""
<html>
  <body>
    <div style="font-size: 12px; color: #666; margin-bottom: 12px;">
      Source: <a href="{source_url}">{source_url}</a>
    </div>
    {html}
  </body>
</html>
""".strip()

    def _generate_pdf_reportlab(
        self,
        filtered_html: str,
        source_url: str,
        images: list[tuple[str, bytes, str, str, str]],
    ) -> bytes:
        """Fallback PDF generator using ReportLab (pure-Python).

        This is not as faithful as WeasyPrint, but keeps Phase 1 usable on Windows
        without installing GTK/Pango dependencies.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.utils import ImageReader
            from reportlab.pdfgen import canvas
        except Exception as e:  # pragma: no cover
            raise RuntimeError(
                "PDF generation requires either WeasyPrint (with native deps) or reportlab. "
                "Install reportlab or switch output_format to 'docx'."
            ) from e

        soup = BeautifulSoup(filtered_html, "lxml")

        # Build a best-effort lookup from original URL -> raw bytes, for non-embedded cases.
        by_url: dict[str, tuple[bytes, str, str]] = {}
        for (_fname, raw, ctype, alt, original_url) in images:
            if original_url:
                by_url[str(original_url)] = (raw, ctype, alt)

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4

        x = 40
        y = height - 50
        line_height = 14

        c.setFont("Helvetica", 10)
        c.drawString(x, y, f"Source: {source_url}")
        y -= line_height * 2

        def new_page() -> None:
            nonlocal y
            c.showPage()
            c.setFont("Helvetica", 9)
            y = height - 50

        def draw_wrapped(text: str, *, font_size: int = 9) -> None:
            nonlocal y
            if not text:
                return
            c.setFont("Helvetica", font_size)
            # Very simple wrapping by characters (keeps it deterministic).
            for raw_line in text.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                while len(line) > 120:
                    chunk, line = line[:120], line[120:]
                    c.drawString(x, y, chunk)
                    y -= line_height
                    if y < 50:
                        new_page()
                c.drawString(x, y, line)
                y -= line_height
                if y < 50:
                    new_page()

        def decode_data_uri(src: str) -> tuple[bytes, str] | None:
            if not src.startswith("data:"):
                return None
            try:
                header, b64 = src.split(",", 1)
                # header: data:<mime>;base64
                mime = header[5:].split(";", 1)[0].strip() or "application/octet-stream"
                return base64.b64decode(b64), mime
            except Exception:
                return None

        def render_image(img_bytes: bytes, ctype: str, *, alt: str, url: str, srcset_url: str = "") -> None:
            nonlocal y
            max_w = width - (2 * x)
            max_h = 260  # points

            # Convert unsupported formats via Pillow if available
            if ctype.lower().startswith("image/webp") or ctype.lower().startswith("image/avif"):
                try:
                    from PIL import Image

                    im = Image.open(io.BytesIO(img_bytes))
                    out = io.BytesIO()
                    im.save(out, format="PNG")
                    img_bytes = out.getvalue()
                    ctype = "image/png"
                except Exception:
                    # If we can't decode, just skip the image.
                    return

            reader = ImageReader(io.BytesIO(img_bytes))
            iw, ih = reader.getSize()
            if not iw or not ih:
                return

            scale = min(max_w / float(iw), max_h / float(ih), 1.0)
            dw = float(iw) * scale
            dh = float(ih) * scale

            # Ensure space (caption + image)
            needed = dh + 12 + (10 if url else 0) + (10 if alt else 0) + 8
            if y - needed < 50:
                new_page()

            # Caption (URL + ALT) first, then image (keeps reading order)
            c.setFont("Helvetica", 7)
            if url:
                c.drawString(x, y, f"Image: {url[:140]}")
                y -= 10
            if srcset_url and srcset_url != url:
                c.drawString(x, y, f"SourceSet: {srcset_url[:140]}")
                y -= 10
            if alt:
                c.drawString(x, y, f"Alt: {alt[:140]}")
                y -= 10
            y -= 2

            c.drawImage(reader, x, y - dh, width=dw, height=dh, preserveAspectRatio=True, mask="auto")
            y -= dh + 14

        def walk(node: Tag) -> None:
            """Walk HTML in-order and render text/images inline (best-effort)."""
            nonlocal y
            for child in list(node.children):
                if isinstance(child, NavigableString):
                    t = str(child).strip()
                    if t:
                        draw_wrapped(t)
                    continue
                if not isinstance(child, Tag):
                    continue

                name = (child.name or "").lower()
                if name == "img":
                    src = (child.get("src") or "").strip()
                    alt = (child.get("alt") or "").strip()
                    orig_src = (child.get("data-ws-img-src") or "").strip()
                    srcset_url = (child.get("data-ws-img-candidate") or "").strip()
                    url = orig_src or (src if src and not src.startswith("data:") else "")

                    decoded = decode_data_uri(src) if src else None
                    if decoded:
                        raw, mime = decoded
                        render_image(raw, mime, alt=alt, url=url, srcset_url=srcset_url)
                        continue

                    if src and src in by_url:
                        raw, ctype, alt2 = by_url[src]
                        render_image(raw, ctype, alt=(alt or alt2), url=url or src, srcset_url=srcset_url)
                        continue

                    # Can't render: leave a placeholder line so the reader knows something was here.
                    draw_wrapped("[image omitted]")
                    continue

                if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    draw_wrapped(child.get_text(" ", strip=True), font_size=11)
                    y -= 4
                    continue
                if name in ("p", "li"):
                    draw_wrapped(child.get_text(" ", strip=True), font_size=9)
                    y -= 2
                    continue
                if name in ("pre", "code"):
                    draw_wrapped(child.get_text("\n", strip=True), font_size=8)
                    y -= 2
                    continue

                # Default: recurse
                walk(child)

        c.setFont("Helvetica", 9)
        root = soup.body if soup.body is not None else soup
        walk(root)

        c.save()
        return buf.getvalue()


